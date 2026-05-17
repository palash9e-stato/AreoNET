"""
Generate Chaukas_Documentation_Updated.docx
Appends new enhancement sections to the existing document.
Uses only styles present in the source docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

SRC = r"c:\Users\Palash\Downloads\CHAUKAS\docc\Chaukas_Documentation.docx"
OUT = r"c:\Users\Palash\Downloads\CHAUKAS\docc\Chaukas_Documentation_Updated.docx"

doc = Document(SRC)

# ── Helpers ─────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    """Manual bullet using '•' prefix since 'List Bullet' style is absent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    run = p.add_run("• " + text)
    run.font.size = Pt(11)

def bold_line(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True

def table_2col(rows, hdr1="Feature", hdr2="Description"):
    t = doc.add_table(rows=1 + len(rows), cols=2)
    # Header row
    hc = t.rows[0].cells
    hc[0].text = hdr1
    hc[1].text = hdr2
    for cell in hc:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    # Data rows
    for i, (c0, c1) in enumerate(rows):
        t.rows[i + 1].cells[0].text = c0
        t.rows[i + 1].cells[1].text = c1
    doc.add_paragraph()

def hr():
    doc.add_paragraph("─" * 90)

# ── Page break ───────────────────────────────────────────────────────────────
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
h1("7. Platform Enhancements — Session 2 (May 2026)")
body(
    "Following the initial submission, the Chaukas platform was significantly extended with "
    "six new modules and a complete UI redesign. These additions directly address real-world "
    "gaps in accident response communication, road infrastructure cost estimation, and "
    "exploratory data transparency."
)
body(f"Documentation generated: {datetime.datetime.now().strftime('%d %B %Y, %H:%M IST')}")

hr()

# ── 7.1 Sidebar UI ───────────────────────────────────────────────────────────
h2("7.1  Professional Dashboard UI Redesign (Sidebar Layout)")
body(
    "The original top navigation bar was replaced with a fixed left sidebar, adopting the "
    "professional SaaS dashboard pattern. This improves navigability across the growing "
    "feature set and presents a more production-ready interface."
)
h3("Design Specifications")
table_2col([
    ("Sidebar Width",        "240 px fixed, collapsible to 64 px icon-only mode"),
    ("Background",           "#0a0f1e (deep navy) with 1 px border-right rgba(255,255,255,0.06)"),
    ("Logo",                 "Red gradient shield icon + CHAUKAS logotype + Road Safety Command subtitle"),
    ("Location Badge",       "Pinned Indore, MP with live green pulse indicator"),
    ("Navigation Sections",  "3 grouped sections: Intelligence / Safety & Features / Community"),
    ("Active State",         "border-left: 2px solid #ef4444 + bg-red-500/10 highlight"),
    ("NEW Badges",           "Cyan 'NEW' pill on all newly added routes"),
    ("User Profile",         "Avatar + role + logout at sidebar bottom"),
    ("3D Globe",             "Retained exclusively on /landing; all inner pages use clean dark background"),
])

h3("Layout Architecture Change")
bullet("App.jsx refactored from absolute-position overlay model to standard flex layout")
bullet("Sidebar: position fixed, z-index 40, rendered for all authenticated inner pages")
bullet("Main content: flex-1, ml-60, overflow-y-auto — pages scroll naturally")
bullet("Landing page retains full-screen 3D globe (React Three Fiber)")
bullet("body overflow:hidden removed to allow sidebar layout scrolling")

hr()

# ── 7.2 Analytics EDA ────────────────────────────────────────────────────────
h2("7.2  Interactive Analytics & EDA Dashboard (/analytics)")
body(
    "The Analytics page was completely rewritten from a static placeholder into a fully "
    "interactive, four-tab Exploratory Data Analysis dashboard. It explains how each "
    "platform module works with live animated visualisations."
)

h3("Tab 1 — Overview")
bullet("4 animated KPI cards with count-up animation (requestAnimationFrame hook)")
bullet("Incident Type Distribution: staggered animated horizontal bar chart")
bullet("Area Risk Summary: clickable table for all 8 Indore neighbourhoods")
bullet("Columns: accidents, potholes, accel (m/s²), severity badge, risk progress bar")

h3("Tab 2 — Hotspot Detection")
bullet("Step-by-step pipeline explainer: Incident Report → AI Severity → DBSCAN → Risk Grid → Heatmap")
bullet("Cluster Risk by Area: horizontal bars coloured by severity tier (critical/high/medium/low)")
bullet("Zone Classification table: Danger ≥75% / Warning ≥50% / Caution ≥25% / Safe <25%")

h3("Tab 3 — Pothole Pipeline")
bullet("6-step pipeline cards with hover-lift animation: Capture → Scoring → Clustering → Severity → Heatmap Fusion → AI Engine")
bullet("Acceleration Intensity Chart: animated bars with Deep/Medium/Shallow classification badges")
bullet("Detection threshold: 14 m/s² | Avg confidence: 92.4%")

h3("Tab 4 — Temporal Trends")
bullet("Hourly Distribution: 24-bar SVG chart with hover tooltips showing incident count + hour")
bullet("Weekly Trend: 7-bar SVG chart — weekend spike (Sat/Sun) clearly visible")
bullet("SVG bars animate from 0% height on mount using delayed setState")
bullet("4 insight cards: Peak Hour (6-7 PM), Quietest (4 AM), Worst Day (Saturday), Weekly Trend (+12%)")

hr()

# ── 7.3 Crash SOS ────────────────────────────────────────────────────────────
h2("7.3  Crash Detection & Multi-Layer Auto-SOS (/crash-sos)")
body(
    "This module closes the most critical gap in accident response: the time between impact "
    "and the arrival of emergency services. The system operates in four simultaneous layers "
    "the moment a crash is confirmed, requiring zero user interaction if the victim is unconscious."
)

h3("Synthetic Accelerometer Data (Always-On Demo)")
body(
    "Because DeviceMotionEvent requires a physical mobile device, a synthetic data generator "
    "was implemented so the feature is always visually demonstrable:"
)
bullet("60 ms setInterval generates continuous road vibration noise: magnitude 2–8 m/s²")
bullet("Every 40–120 ticks (~4–8 seconds), a random pothole spike (14–30 m/s²) is injected")
bullet("Simulation pauses during active crash countdown (crashRef flag)")
bullet("XYZ axis values (x, y, z) animate live on the dashboard cards")
bullet("'Simulate Crash' injects a 58 m/s² spike and launches the full SOS pipeline")

h3("5-Channel Silent SOS Broadcast (Layer 1)")
table_2col([
    ("Indore CATS Ambulance (108)", "Notified at t+500 ms"),
    ("Indore Police Control (100)", "Notified at t+1200 ms"),
    ("Fire Brigade (101)",          "Notified at t+1800 ms"),
    ("Community Mesh (3 nearby)",   "Good Samaritan network alerted at t+400 ms"),
    ("Personal Emergency Contacts", "Pre-registered contacts notified at t+300 ms"),
], "Channel", "Dispatch Timing")
body("Each dispatch is signed with GPS coordinate + device fingerprint + timestamp for tamper-proof evidence.")

h3("Good Samaritan Network (Layer 1 — Community)")
bullet("Three nearest Chaukas community members within 2 km radius are alerted simultaneously")
bullet("Response cards appear on dispatch screen showing name and distance")
bullet("Extends conventional emergency dispatch with a civilian first-responder network")

h3("Auto Scene Documentation (Layer 3 — Feature 4)")
bullet("Phase 1 — Capturing (0–1.5 s): 3 auto-photos (front, back, left) + 5-second video clip")
bullet("Phase 2 — Uploading (1.5–3.5 s): GPS-tagged assets uploaded to Chaukas server")
bullet("Phase 3 — Done: Crash Report PDF generated with timestamp, coordinates, media inventory")
bullet("If Medical Card exists in localStorage, it is automatically embedded in the report")
bullet("Download PDF button surfaces the report for FIR filing and insurance submission")

h3("Offline SMS Fallback (Layer 4 — Feature 5)")
bullet("If navigator.onLine returns false, platform switches to SMS mode automatically")
bullet("sms:100 and sms:108 triggered via window.open('sms:NUMBER?body=...', '_blank')")
bullet("Pre-composed message includes Google Maps deep-link with exact GPS coordinates")
bullet("Timestamp appended for legal evidence value in FIR")
bullet("Manual 'Send SMS Now' button always available from post-SOS view")

hr()

# ── 7.4 Medical Card ──────────────────────────────────────────────────────────
h2("7.4  Smart Medical Emergency Card (/medical-card) — Feature 3")
body(
    "Paramedics arriving at crash scenes often lack critical patient information that can "
    "determine life-or-death treatment decisions within the first minutes. The Medical Card "
    "provides a zero-friction, zero-login solution that auto-attaches to every SOS dispatch."
)

h3("Data Captured")
table_2col([
    ("Full Name & Age",         "Identity fields"),
    ("Blood Type",              "A+ / A- / B+ / B- / AB+ / AB- / O+ / O- selector"),
    ("Allergies",               "Free-text: e.g. Penicillin, Peanuts, Latex"),
    ("Pre-existing Conditions", "e.g. Diabetes, Hypertension, Epilepsy"),
    ("Current Medications",     "e.g. Metformin 500mg, Aspirin 75mg"),
    ("Emergency Contact",       "Name + phone number"),
])

h3("Secure Token URL System")
bullet("On save, a cryptographic token (random 8-char alphanumeric) is generated client-side")
bullet("Secure URL format: https://chaukas.in/sos/{TOKEN}")
bullet("URL is accessible without login — paramedic opens it in 1 tap from dispatch message")
bullet("Designed to auto-expire after 2 hours to protect patient privacy")
bullet("Card data persisted in localStorage — zero cloud dependency")
bullet("Token is automatically embedded in every Crash SOS dispatch packet at fire-time")

h3("Paramedic Card Preview")
bullet("Live preview renders exactly what the paramedic sees: large blood type badge, colour-coded rows")
bullet("Sections: Allergies (yellow), Conditions (red), Medications (purple), Emergency Contact (green)")
bullet("URGENT badge with animate-pulse effect for immediate paramedic attention")
bullet("Preview updates in real-time as user edits the form")

hr()

# ── 7.5 Night Watch ───────────────────────────────────────────────────────────
h2("7.5  Night Watch — Convoy Safety Mode (/night-watch)")
body(
    "Addresses the scenario where a driver breaks down or is stranded inside a known accident "
    "hotspot at night. Night Watch passively monitors for dangerous stops and auto-alerts "
    "emergency contacts if the user does not respond."
)

h3("Core Detection Logic")
bullet("User activates Night Watch toggle — GPS polling begins")
bullet("System checks every second if GPS coordinates fall within a registered Danger Zone")
bullet("If stopped for more than 2 minutes (120 seconds) inside a Danger Zone: auto-alert fires")
bullet("Orange progress bar visualises the 2-minute countdown in real time")
bullet("10-second cancel window with 'I'm Safe' button before alert is dispatched")

h3("Monitored Danger Zones — Indore")
table_2col([
    ("Rajwada Chowk",  "22.7196, 75.8577 — Critical"),
    ("Palasia",        "22.7245, 75.8400 — Critical"),
    ("LIG Square",     "22.7050, 75.8650 — High"),
    ("Bhawarkuan",     "22.7155, 75.8800 — High"),
], "Zone", "Coordinates & Risk Level")

h3("Emergency Contact Management")
bullet("Dynamic add/remove contacts UI (name + phone number)")
bullet("Alert packet includes: GPS coordinates, zone name, timestamp, user identity")
bullet("Dismissable success state with 'Safe check-in confirmed' confirmation")

hr()

# ── 7.6 Repair Estimator ──────────────────────────────────────────────────────
h2("7.6  Pothole Repair Cost Estimator (/repair-estimator)")
body(
    "Municipal engineers currently survey potholes manually and estimate material requirements "
    "by hand, causing delays and under-procurement. The Repair Estimator automates this "
    "calculation directly from the platform's pothole detection data."
)

h3("Material Volume Formula")
body("Volume per pothole (litres) = (Length cm × Width cm × Depth cm) / 1000")
table_2col([
    ("Deep Pothole",    "L=50cm, W=40cm, D=12cm — 24.0 litres per pothole"),
    ("Medium Pothole",  "L=40cm, W=35cm, D=6cm  —  8.4 litres per pothole"),
    ("Shallow Pothole", "L=30cm, W=25cm, D=3cm  —  2.25 litres per pothole"),
    ("HMA Density",     "2.4 kg/litre (Hot Mix Asphalt industry standard)"),
    ("Material Cost",   "₹8/kg (PWD rate schedule 2024)"),
    ("Labour",          "₹150 per pothole (standard fill and compact)"),
    ("Traffic Mgmt",    "₹1,200 flat rate per area repair visit"),
], "Parameter", "Value / Rate")

h3("Per-Area Output")
bullet("Total volume (litres) and Hot Mix Asphalt required (kg)")
bullet("Itemised Bill of Materials: material cost + labour + traffic management")
bullet("Grand total estimate per area with priority flag (URGENT / HIGH / NORMAL)")
bullet("City-wide aggregate: total potholes, total HMA kg, total estimated budget for Indore")
bullet("Sortable table — top 4 rows shown by default, 'Show all' expands to all 8 areas")

h3("Indore Area Coverage")
body(
    "8 neighbourhoods: Rajwada, Palasia, Bhawarkuan, Vijay Nagar, LIG Square, "
    "Khajrana, Scheme 54, Super Corridor."
)

hr()

# ── 7.7 Summary ───────────────────────────────────────────────────────────────
h2("7.7  New Module Summary")
table_2col([
    ("/analytics",        "4-tab interactive EDA: KPIs, hotspot pipeline, pothole pipeline, temporal trends"),
    ("/crash-sos",        "5-layer silent broadcast, scene auto-capture, PDF report, offline SMS fallback, synthetic accel chart"),
    ("/medical-card",     "Secure paramedic card with token URL, no-login access, auto-attached to SOS dispatch"),
    ("/night-watch",      "2-minute danger-zone stop monitor, emergency contact alerts, dynamic contact management"),
    ("/repair-estimator", "HMA volume formula, itemised BoM, labour + traffic cost, city-wide budget dashboard"),
    ("Sidebar UI",        "Fixed left sidebar, 3 nav groups, collapsible mode, NEW badges, live Indore status"),
], "Route / Module", "Key Capabilities")

hr()

# ── 7.8 Revised Architecture ──────────────────────────────────────────────────
h2("7.8  Revised System Architecture")
body("The platform now follows a sidebar-first layout with clean separation:")
body(
    "BROWSER\n"
    "  /landing            Full-screen 3D Globe (React Three Fiber) + LandingPage\n"
    "  /login              Full-screen Login form\n"
    "  /[all other routes] Fixed Sidebar + Scrollable Page Content\n"
    "        /intelligence     Crisis Incident Dashboard\n"
    "        /analytics        Interactive EDA Dashboard [NEW]\n"
    "        /hotspot          ML Hotspot + Pothole Heatmap\n"
    "        /crash-sos        Crash Detection & SOS Dispatch [NEW]\n"
    "        /medical-card     Medical Emergency Card [NEW]\n"
    "        /night-watch      Night Watch Convoy Safety [NEW]\n"
    "        /repair-estimator Pothole Repair Cost Estimator [NEW]\n"
    "        /severity         AI Severity Engine\n"
    "        /escalation       NDMA State Machine Escalation\n"
    "        /community        Community Hub\n"
    "        /docs             Documentation\n"
)

hr()

# ── 7.9 Communication Gap Analysis ────────────────────────────────────────────
h2("7.9  Communication Gap Analysis — Before vs. After")
table_2col([
    ("Crash to first responder notified",  "Before: 8–15 min (manual dial)      | After: <5 seconds (auto-broadcast)"),
    ("Medical info to paramedic",          "Before: Not available               | After: Instant via Medical Card token URL"),
    ("Scene documentation",                "Before: Manual, often forgotten      | After: Auto-captured + GPS-tagged PDF"),
    ("No-internet scenario",               "Before: Total communication loss     | After: SMS fallback (sms: protocol)"),
    ("Isolated breakdown detection",       "Before: No mechanism                 | After: Night Watch 2-min stop trigger"),
    ("Pothole repair procurement",         "Before: Manual survey, weeks delay   | After: Auto BoM + itemised cost estimate"),
], "Communication Gap", "Resolution")

body("")
body(
    "All enhancements are production-demo ready and fully integrated into the live platform "
    "running at http://localhost:5173. The platform requires no external APIs for these new "
    "features — all data is processed client-side or via the local FastAPI backend."
)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nSaved successfully: {OUT}")
